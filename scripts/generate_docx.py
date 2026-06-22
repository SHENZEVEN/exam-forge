#!/usr/bin/env python3
"""将 Markdown 模拟卷转换为 Word 文档。

自动添加页眉、分页符、姓名/学号填空区。
纯 python-docx 实现，不含外部模板依赖。

用法:
    python3 generate_docx.py 模拟卷.md -o 模拟卷.docx --course "数据结构"

依赖:
    pip install python-docx
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("请先安装 python-docx：pip install python-docx")

# ── 配置 ──────────────────────────────────────────────
FONT_BODY = "宋体"          # 正文中文
FONT_BODY_EN = "Times New Roman"  # 正文英文/数字
FONT_SIZE_BODY = 12         # 正文字号（pt）
FONT_SIZE_TITLE = 16        # 标题字号
FONT_SIZE_SECTION = 14      # 题型标题字号
SEPARATOR_MARKER = "══════"  # 试卷与答案之间的分隔标记


# ── 工具函数 ──────────────────────────────────────────
def set_font(run, name_cn=FONT_BODY, name_en=FONT_BODY_EN, size=FONT_SIZE_BODY, bold=False):
    """给 run 设置中英文字体（解决中文显示问题）。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name_en
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), name_cn)


def add_page_break(doc):
    """插入分页符。"""
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p._element.append(OxmlElement("w:br"))
    p._element[-1].set(qn("w:type"), "page")


def add_header_text(doc, course_name):
    """给文档每节添加页眉。"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{course_name}  模拟试卷（一）")
        set_font(run, size=9)


def add_student_info(doc):
    """在试卷开头添加姓名/学号填空区。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("姓名：__________    学号：__________    得分：__________")
    set_font(run)
    doc.add_paragraph()  # 空行


def parse_and_build(doc, md_text, course_name):
    """解析 Markdown 文本，逐段写入 docx。

    处理：
    - # 标题 → 居中大标题
    - ## 标题 → 左对齐节标题
    - ═══ 分隔线 → 分页符 + "参考答案与解析" 区域
    - 普通段落 → 正文
    - 表格行（|...|...|） → 保留为等宽文本
    """
    lines = md_text.split("\n")
    after_separator = False
    added_student_info = False

    for line in lines:
        stripped = line.strip()

        # ── 分隔标记：试卷部分结束，进入答案部分 ──
        if SEPARATOR_MARKER in stripped:
            add_page_break(doc)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("📋 参考答案与解析")
            set_font(run, size=FONT_SIZE_SECTION, bold=True)
            doc.add_paragraph()
            after_separator = True
            continue

        # ── 一级标题：试卷名称 ──
        if stripped.startswith("# ") and not stripped.startswith("## "):
            title = stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            set_font(run, size=FONT_SIZE_TITLE, bold=True)
            doc.add_paragraph()
            # 标题后添加学生信息填空区
            if not added_student_info and not after_separator:
                add_student_info(doc)
                added_student_info = True
            continue

        # ── 二级标题：题型名称 ──
        if stripped.startswith("## "):
            if not added_student_info and not after_separator:
                add_student_info(doc)
                added_student_info = True
            section_title = stripped[3:]
            p = doc.add_paragraph()
            run = p.add_run(section_title)
            set_font(run, size=FONT_SIZE_SECTION, bold=True)
            doc.add_paragraph()
            continue

        # ── 空行 ──
        if not stripped:
            continue

        # ── 普通段落 ──
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        set_font(run)


# ── 主入口 ────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="将 Markdown 模拟卷转换为 Word 文档"
    )
    parser.add_argument("input", help="输入的 Markdown 文件路径")
    parser.add_argument("-o", "--output", required=True, help="输出的 .docx 文件路径")
    parser.add_argument("--course", default="", help="课程名称（用于页眉）")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        sys.exit(f"文件不存在: {args.input}")

    md_text = input_path.read_text(encoding="utf-8")
    course_name = args.course or "课程名称"

    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 添加页眉
    add_header_text(doc, course_name)

    # 解析 Markdown 并写入
    parse_and_build(doc, md_text, course_name)

    # 保存
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✅ 已导出: {output_path}")


if __name__ == "__main__":
    main()
